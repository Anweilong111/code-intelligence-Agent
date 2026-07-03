# -*- coding: utf-8 -*-
"""Build the beginner-friendly project report as a polished DOCX.

Source: PROJECT_REPORT_BEGINNER_GUIDE.md
Output: Code_Intelligence_Agent_Project_Report.docx
"""

from __future__ import annotations

import importlib.util
import os
import re
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "PROJECT_REPORT_BEGINNER_GUIDE.md"
OUTPUT = ROOT / "Code_Intelligence_Agent_Project_Report.docx"
SKILL_TABLE_GEOMETRY = Path(
    os.environ.get("DOCUMENT_TABLE_GEOMETRY", "")
) if "DOCUMENT_TABLE_GEOMETRY" in os.environ else Path()

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F4F6F9"
WHITE = "FFFFFF"


def load_table_geometry():
    if SKILL_TABLE_GEOMETRY.exists():
        spec = importlib.util.spec_from_file_location("table_geometry", SKILL_TABLE_GEOMETRY)
        module = importlib.util.module_from_spec(spec)
        assert spec and spec.loader
        spec.loader.exec_module(module)
        return module.apply_table_geometry
    return apply_table_geometry_fallback


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def apply_table_geometry_fallback(
    table,
    column_widths_dxa: Sequence[int],
    *,
    table_width_dxa: int | None = None,
    indent_dxa: int | None = None,
    cell_margins_dxa: dict[str, int] | None = None,
) -> None:
    widths = [int(width) for width in column_widths_dxa]
    width_total = int(table_width_dxa if table_width_dxa is not None else sum(widths))
    margins = dict(CELL_MARGINS_DXA)
    if cell_margins_dxa:
        margins.update(cell_margins_dxa)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = _ensure_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_total))
    tbl_ind = _ensure_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA if indent_dxa is None else indent_dxa))
    layout = _ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = Twips(width)
    for row in table.rows:
        row.height = None
        for col_idx, cell in enumerate(row.cells):
            width = widths[col_idx]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = _ensure_child(tc_pr, "w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            tc_mar = _ensure_child(tc_pr, "w:tcMar")
            for side in ("top", "bottom", "start", "end"):
                margin = _ensure_child(tc_mar, f"w:{side}")
                margin.set(qn("w:type"), "dxa")
                margin.set(qn("w:w"), str(margins[side]))


APPLY_TABLE_GEOMETRY = load_table_geometry()


def set_font_name(run_or_font, ascii_font: str, east_asia_font: str | None = None) -> None:
    east_asia_font = east_asia_font or ascii_font
    if hasattr(run_or_font, "font"):
        run_or_font.font.name = ascii_font
        r_pr = run_or_font._element.get_or_add_rPr()
    else:
        run_or_font.name = ascii_font
        r_pr = run_or_font._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def set_style_font(style, ascii_font: str, east_asia_font: str, size_pt: float, color=None, bold=None) -> None:
    font = style.font
    font.name = ascii_font
    font.size = Pt(size_pt)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = _ensure_child(tc_pr, "w:shd")
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = _ensure_child(p_pr, "w:shd")
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font_name(run, "Calibri", "Microsoft YaHei")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", "Microsoft YaHei", 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    set_style_font(h1, "Calibri", "Microsoft YaHei", 16, BLUE, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, "Calibri", "Microsoft YaHei", 13, BLUE, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, "Calibri", "Microsoft YaHei", 12, DARK_BLUE, True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        set_style_font(style, "Calibri", "Microsoft YaHei", 11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", 1)
    else:
        code_style = styles["CodeBlock"]
    set_style_font(code_style, "Consolas", "Microsoft YaHei", 9.5, RGBColor(0x22, 0x22, 0x22))
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.line_spacing = 1.05

    if "TOCMain" not in styles:
        toc_main = styles.add_style("TOCMain", 1)
    else:
        toc_main = styles["TOCMain"]
    set_style_font(toc_main, "Calibri", "Microsoft YaHei", 10.5, NAVY, True)
    toc_main.paragraph_format.space_after = Pt(2)

    if "TOCSub" not in styles:
        toc_sub = styles.add_style("TOCSub", 1)
    else:
        toc_sub = styles["TOCSub"]
    set_style_font(toc_sub, "Calibri", "Microsoft YaHei", 9.5, GRAY)
    toc_sub.paragraph_format.left_indent = Inches(0.22)
    toc_sub.paragraph_format.space_after = Pt(1)

    header = section.header.paragraphs[0]
    header.text = "Code Intelligence Agent System | 项目完整报告与学习指南"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_font_name(run, "Calibri", "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def clean_inline(text: str) -> str:
    return text.replace("  ", " ").strip()


TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_markdown_runs(paragraph, text: str, *, base_size: float | None = None, bold_all: bool = False) -> None:
    for token in TOKEN_RE.split(text):
        if not token:
            continue
        run = paragraph.add_run()
        if token.startswith("**") and token.endswith("**"):
            run.text = token[2:-2]
            run.bold = True
            set_font_name(run, "Calibri", "Microsoft YaHei")
        elif token.startswith("`") and token.endswith("`"):
            run.text = token[1:-1]
            set_font_name(run, "Consolas", "Microsoft YaHei")
            run.font.size = Pt(9.5 if base_size is None else base_size)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        else:
            run.text = token
            set_font_name(run, "Calibri", "Microsoft YaHei")
        if base_size is not None and not (token.startswith("`") and token.endswith("`")):
            run.font.size = Pt(base_size)
        if bold_all:
            run.bold = True


def add_para(doc: Document, text: str, style: str | None = None, *, align=None, bold_all=False, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    add_markdown_runs(p, clean_inline(text), base_size=size, bold_all=bold_all)
    return p


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(42)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("项目完整报告 / 初学者学习指南")
    set_font_name(run, "Calibri", "Microsoft YaHei")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = BLUE

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Code Intelligence Agent System")
    set_font_name(run, "Calibri", "Microsoft YaHei")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("基于程序分析与大语言模型的代码智能体系统")
    set_font_name(run, "Calibri", "Microsoft YaHei")
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_BLUE

    meta_rows = [
        ("文档定位", "完整项目报告、从 0 学习路线、简历包装与面试问答"),
        ("项目方向", "程序分析 / 缺陷定位 / 自动修复 / LLM Agent / 搜索与评估"),
        ("核心闭环", "Repo Understanding -> Fault Localization -> Patch Generation -> Sandbox Validation -> Reflection Loop"),
        ("生成日期", date.today().isoformat()),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(meta_rows):
        cells = table.rows[row_idx].cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE_GRAY)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font_name(r, "Calibri", "Microsoft YaHei")
                    r.font.size = Pt(10.5)
        for r in cells[0].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = NAVY
    APPLY_TABLE_GEOMETRY(table, [2100, 7260], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa=CELL_MARGINS_DXA)

    doc.add_paragraph()
    add_callout(
        doc,
        [
            "简历表达重点：不要写成“调用大模型修代码”，要写成“程序图建模 + 缺陷定位排序 + 多候选补丁搜索 + 执行验证闭环”。",
            "当前项目适合写算法向 Agent 项目；后续深化可继续做任意 GitHub repo 自动 benchmark 化和 Learning-to-Rank 缺陷定位。",
        ],
        fill="F4F6F9",
    )
    doc.add_page_break()


def collect_headings(lines: Iterable[str]) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for raw in lines:
        line = raw.strip()
        match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if match:
            level = len(match.group(1)) - 1
            text = re.sub(r"`([^`]+)`", r"\1", match.group(2)).strip()
            headings.append((level, text))
    return headings


def add_static_toc(doc: Document, headings: Sequence[tuple[int, str]]) -> None:
    doc.add_heading("目录", level=1)
    lead = add_para(doc, "本目录为静态目录，用于快速了解报告结构；具体页码以 Word 打开后的分页为准。")
    lead.runs[0].font.color.rgb = GRAY
    for level, text in headings:
        if level == 1:
            add_para(doc, text, "TOCMain")
    doc.add_page_break()


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def compute_widths(num_cols: int) -> list[int]:
    if num_cols == 1:
        return [CONTENT_WIDTH_DXA]
    if num_cols == 2:
        return [2700, CONTENT_WIDTH_DXA - 2700]
    if num_cols == 3:
        return [2200, 3580, CONTENT_WIDTH_DXA - 5780]
    if num_cols == 4:
        return [1800, 2500, 2500, CONTENT_WIDTH_DXA - 6800]
    base = CONTENT_WIDTH_DXA // num_cols
    widths = [base] * num_cols
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def style_cell_text(cell, *, header: bool = False) -> None:
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_font_name(run, "Calibri", "Microsoft YaHei")
            run.font.size = Pt(9.5)
            if header:
                run.bold = True
                run.font.color.rgb = NAVY


def add_markdown_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(num_cols):
            cell = table.rows[r_idx].cells[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = text.replace("<br>", "\n")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE_GRAY)
            else:
                set_cell_shading(cell, WHITE)
            style_cell_text(cell, header=(r_idx == 0))
    APPLY_TABLE_GEOMETRY(table, compute_widths(num_cols), table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa=CELL_MARGINS_DXA)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    if not lines:
        return
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = "CodeBlock"
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_font_name(run, "Consolas", "Microsoft YaHei")
        run.font.size = Pt(9.2)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    APPLY_TABLE_GEOMETRY(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120})
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(doc: Document, lines: Sequence[str], *, fill: str = LIGHT_GRAY) -> None:
    if not lines:
        return
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_markdown_runs(p, clean_inline(line), base_size=10.5)
        for run in p.runs:
            run.font.color.rgb = NAVY
    APPLY_TABLE_GEOMETRY(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120})
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_markdown_image(doc: Document, alt_text: str, image_path: str) -> None:
    path = Path(image_path)
    if not path.is_absolute():
        path = ROOT / path
    if not path.exists():
        add_callout(
            doc,
            [f"图示未找到：{alt_text}（{image_path}）"],
            fill="FFF7E6",
        )
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    if alt_text:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)
        caption_run = caption.add_run(alt_text)
        set_font_name(caption_run, "Calibri", "Microsoft YaHei")
        caption_run.font.size = Pt(9.2)
        caption_run.italic = True
        caption_run.font.color.rgb = GRAY


def markdown_to_docx(doc: Document, lines: list[str]) -> None:
    idx = 0
    in_code = False
    code_lines: list[str] = []
    skip_first_h1 = True

    while idx < len(lines):
        raw = lines[idx].rstrip("\n")
        line = raw.strip()

        if in_code:
            if line.startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(raw)
            idx += 1
            continue

        if line.startswith("```"):
            in_code = True
            code_lines = []
            idx += 1
            continue

        if not line:
            idx += 1
            continue

        if line == "---":
            idx += 1
            continue

        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line)
        if image:
            add_markdown_image(doc, image.group(1).strip(), image.group(2).strip())
            idx += 1
            continue

        if line.startswith("|") and idx + 1 < len(lines) and is_table_separator(lines[idx + 1].strip()):
            table_rows = [parse_table_row(line)]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                if not is_table_separator(lines[idx].strip()):
                    table_rows.append(parse_table_row(lines[idx].strip()))
                idx += 1
            add_markdown_table(doc, table_rows)
            continue

        if line.startswith(">"):
            quote_lines: list[str] = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote_lines.append(lines[idx].strip().lstrip(">").strip())
                idx += 1
            add_callout(doc, quote_lines)
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                idx += 1
                continue
            word_level = min(max(level - 1, 1), 3)
            doc.add_heading(re.sub(r"`([^`]+)`", r"\1", text), level=word_level)
            idx += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            add_para(doc, bullet.group(1), "List Bullet")
            idx += 1
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered:
            add_para(doc, numbered.group(1), "List Number")
            idx += 1
            continue

        add_para(doc, line)
        idx += 1

    if in_code and code_lines:
        add_code_block(doc, code_lines)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Missing source Markdown: {SOURCE}")
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_static_toc(doc, collect_headings(lines))
    markdown_to_docx(doc, lines)

    doc.core_properties.title = "Code Intelligence Agent System 项目完整报告与学习指南"
    doc.core_properties.subject = "程序分析、缺陷定位、自动修复、LLM Agent、简历与面试"
    doc.core_properties.author = "Code Intelligence Agent"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
