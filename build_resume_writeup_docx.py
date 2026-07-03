# -*- coding: utf-8 -*-
"""Build a resume-focused Word document for the Code Intelligence Agent project."""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Code_Intelligence_Agent_Resume_Writeup.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 90, "bottom": 90, "start": 130, "end": 130}

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
GRAY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x00, 0x00, 0x00)
LIGHT_BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
SOFT_BLUE = "F4F8FC"
SOFT_GREEN = "EEF7F0"
SOFT_YELLOW = "FFF7E6"
WHITE = "FFFFFF"


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_font(run_or_style_font, ascii_font: str = "Calibri", east_asia_font: str = "Microsoft YaHei") -> None:
    if hasattr(run_or_style_font, "font"):
        run_or_style_font.font.name = ascii_font
        r_pr = run_or_style_font._element.get_or_add_rPr()
    else:
        run_or_style_font.name = ascii_font
        r_pr = run_or_style_font._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def set_style(style, size: float, color=None, bold: bool | None = None) -> None:
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    set_font(style.font)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = _ensure_child(tc_pr, "w:shd")
    shd.set(qn("w:fill"), fill)


def apply_table_geometry(
    table,
    column_widths_dxa: Sequence[int],
    *,
    table_width_dxa: int = CONTENT_WIDTH_DXA,
    indent_dxa: int = TABLE_INDENT_DXA,
    cell_margins_dxa: dict[str, int] | None = None,
) -> None:
    widths = [int(width) for width in column_widths_dxa]
    margins = dict(CELL_MARGINS_DXA)
    if cell_margins_dxa:
        margins.update(cell_margins_dxa)

    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = _ensure_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(table_width_dxa))
    tbl_ind = _ensure_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    layout = _ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = Twips(width)
    for row in table.rows:
        row.height = None
        for col_idx, cell in enumerate(row.cells):
            width = widths[col_idx]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = _ensure_child(tc_pr, "w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            tc_mar = _ensure_child(tc_pr, "w:tcMar")
            for side in ("top", "bottom", "start", "end"):
                margin = _ensure_child(tc_mar, f"w:{side}")
                margin.set(qn("w:type"), "dxa")
                margin.set(qn("w:w"), str(margins[side]))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style(styles["Normal"], 11, BLACK)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    set_style(styles["Heading 1"], 16, BLUE, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style(styles["Heading 2"], 13, BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(7)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style(styles["Heading 3"], 12, DARK_BLUE, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        set_style(style, 10.6, BLACK)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = section.header.paragraphs[0]
    header.text = "Code Intelligence Agent System | 简历写法与面试表达"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY


def add_run(paragraph, text: str, *, size: float = 11, bold: bool = False, color=None, font: str = "Calibri"):
    run = paragraph.add_run(text)
    set_font(run, font, "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def add_paragraph(doc: Document, text: str, *, size: float = 11, bold: bool = False, color=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, size=size, bold=bold, color=color)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    add_run(p, text, size=10.6)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    add_run(p, text, size=10.6)
    return p


def style_cell(cell, *, header: bool = False, size: float = 9.5) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_font(run)
            run.font.size = Pt(size)
            if header:
                run.bold = True
                run.font.color.rgb = NAVY


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE_GRAY)
        style_cell(cell, header=True)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            set_cell_shading(cell, WHITE)
            style_cell(cell)

    apply_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_box(doc: Document, lines: Sequence[str], *, title: str | None = None, fill: str = SOFT_BLUE, monospace: bool = False):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    cell.text = ""
    if title:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        add_run(p, title, size=10.6, bold=True, color=NAVY)
        start_idx = 1
    else:
        start_idx = 0
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 and not title else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_run(
            p,
            line,
            size=9.7 if monospace else 10.3,
            color=BLACK,
            font="Consolas" if monospace else "Calibri",
        )
    apply_table_geometry(
        table,
        [CONTENT_WIDTH_DXA],
        cell_margins_dxa={"top": 110, "bottom": 110, "start": 150, "end": 150},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(30)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(kicker, "简历写法 / 面试表达 / 项目边界", size=12, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    add_run(title, "Code Intelligence Agent System", size=26, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    add_run(subtitle, "如何把算法向代码智能体项目写进简历", size=15, color=DARK_BLUE)

    add_table(
        doc,
        ["项目", "内容"],
        [
            ["项目定位", "程序分析 + 缺陷定位 + 补丁搜索 + sandbox 执行验证的算法向 LLM Agent 项目"],
            ["适合岗位", "算法工程师、LLM Agent 工程师、代码智能/AI Infra/开发工具方向"],
            ["核心亮点", "Program Graph、FinalScore、Beam Search、Reflection Loop、Ablation Study"],
            ["当前边界", "已支持受控 GitHub raw-source mutation benchmark；不应表述为任意 GitHub repo 一键真实修复并提交 PR"],
            ["生成日期", date.today().isoformat()],
        ],
        [2100, 7260],
    )
    add_box(
        doc,
        [
            "简历表达原则：不要写成“调用大模型修代码”，要写成“程序图建模 + 多信号缺陷定位 + 候选补丁搜索 + 执行验证闭环”。",
            "面试表达原则：指标要放在受控 benchmark/template 的上下文里说明，既突出算法深度，也避免夸大真实世界泛化能力。",
        ],
        fill=SOFT_GREEN,
    )
    doc.add_page_break()


def add_final_resume_version(doc: Document) -> None:
    doc.add_heading("1. 最终推荐写法", level=1)
    add_paragraph(doc, "这部分可以直接放进简历。建议项目名称保留技术密度，项目 bullet 控制在 3 条左右。")

    add_box(
        doc,
        [
            "Code Intelligence Agent System：基于程序分析与 LLM 的代码缺陷定位与自动修复 Agent",
            "",
            "项目描述：构建面向 Python 仓库的代码智能体系统，融合 AST/CFG/Call Graph/Data-flow 程序分析、FinalScore 缺陷定位、Beam Search 补丁搜索、pytest sandbox 执行验证与 ablation study，实现从代码理解、Bug 定位到自动修复的可复现实验闭环。",
        ],
        title="项目名称与一句话描述",
        fill=SOFT_BLUE,
    )

    add_table(
        doc,
        ["简历 Bullet", "为什么这样写"],
        [
            [
                "构建基于 AST/CFG/Call Graph/Data-flow 的代码智能体，支持 Python 仓库解析、静态规则缺陷检测、函数级缺陷定位与 Top-k suspicious ranking，在 62-case GitHub raw-source mutation benchmark 上实现 Top-1/Top-3 Localization 均为 1.0000。",
                "突出程序分析和定位指标，说明不是简单 LLM wrapper。",
            ],
            [
                "设计融合 SBFL、StaticRuleScore、GraphScore、SemanticScore 与 RiskScore 的 FinalScore 排序框架，并输出 attribution、Top-1 margin 与 counterfactual ranking 证据，增强定位结果的可解释性。",
                "突出算法设计、可解释性和排序建模能力。",
            ],
            [
                "实现 sandbox-validated 自动修复闭环，支持 rule/LLM patch generation、patch validation、Beam Search、execution feedback、reflection repair、candidate deduplication 与 diversity reranking，在受控 benchmark 上 Patch Success Rate 达 1.0000，Beam Success Rate 达 0.9516。",
                "突出搜索、执行验证、自我修复和工程闭环。",
            ],
            [
                "搭建 benchmark runner、quality gate、ablation study、hard-case mining/generation 与 showcase report，自动输出定位指标、补丁搜索轨迹、消融贡献和简历展示报告，用实验证据验证各算法模块贡献。",
                "突出评估体系，这是算法向项目比普通工程项目更有说服力的部分。",
            ],
        ],
        [4800, 4560],
    )

    add_box(
        doc,
        [
            "技术栈：Python, AST, CFG, Call Graph, Data-flow Analysis, Program Graph, Program Slicing, SBFL, Fault Localization, Beam Search, Patch Generation, pytest Sandbox, LLM Agent, Ablation Study",
        ],
        title="技术栈写法",
        fill=SOFT_BLUE,
    )


def add_versions(doc: Document) -> None:
    doc.add_heading("2. 不同简历空间下的版本", level=1)

    add_box(
        doc,
        [
            "实现基于程序图与 LLM 的代码自动修复 Agent，融合 AST/CFG/Call Graph/Data-flow、FinalScore 缺陷定位、Beam Search 补丁搜索和 pytest sandbox 验证，在 62-case 受控 benchmark 上完成定位与修复闭环评估。",
        ],
        title="极简版：只放 1 条 bullet",
        fill=SOFT_BLUE,
    )

    add_table(
        doc,
        ["版本", "适用场景", "推荐内容"],
        [
            [
                "2 条 bullet",
                "简历项目很多，需要压缩篇幅",
                "第一条写程序图建模 + FinalScore 定位；第二条写 Beam Search + sandbox 修复闭环 + benchmark 指标。",
            ],
            [
                "3 条 bullet",
                "最推荐",
                "第一条写仓库理解与程序图；第二条写定位排序算法；第三条写修复搜索与评估体系。",
            ],
            [
                "4 条 bullet",
                "投算法/Agent/代码智能方向",
                "额外增加 ablation、hard-case generation、quality gate，体现实验设计和算法贡献证明。",
            ],
        ],
        [1600, 2900, 4860],
    )

    add_box(
        doc,
        [
            "1. 构建 AST/CFG/Call Graph/Data-flow 驱动的 Program Graph，完成 Python 仓库理解、静态规则检测和函数级 Top-k 缺陷定位。",
            "2. 设计 FinalScore 多信号排序算法，融合 SBFL、静态规则、图传播、语义相似度和风险信号，并输出 attribution 解释。",
            "3. 实现 Beam Search + pytest sandbox 的自动修复闭环，结合执行反馈与 reflection loop 迭代补丁，并通过 benchmark、quality gate 和 ablation study 验证效果。",
        ],
        title="最稳妥的 3 条 bullet 版本",
        fill=SOFT_GREEN,
    )


def add_role_mapping(doc: Document) -> None:
    doc.add_heading("3. 按岗位方向调整写法", level=1)
    add_table(
        doc,
        ["投递方向", "重点强调", "简历措辞"],
        [
            [
                "算法工程师",
                "排序、图建模、SBFL、消融实验",
                "设计多信号缺陷定位排序算法 FinalScore，融合图结构、测试覆盖、静态规则和语义信号，并通过 ablation study 验证各信号贡献。",
            ],
            [
                "LLM Agent 工程师",
                "工具调用、反思闭环、执行反馈",
                "实现面向代码修复的 Agent loop，将 patch generation、sandbox execution、failure classification 与 reflection repair 串成可验证闭环。",
            ],
            [
                "AI Infra/开发工具",
                "可复现、质量门禁、自动报告",
                "搭建 benchmark runner、quality gate 与 showcase report，沉淀可复现实验产物和自动化评估流程。",
            ],
            [
                "软件工程/平台",
                "模块化、测试、工程可靠性",
                "按 repo parser、program graph、localizer、repair loop、search、evaluation 分层实现，并用 pytest 回归保障模块行为。",
            ],
        ],
        [1700, 2500, 5160],
    )


def add_explanation(doc: Document) -> None:
    doc.add_heading("4. 为什么这个项目适合写成算法向 Agent", level=1)
    add_paragraph(
        doc,
        "它的价值不在于“接了一个大模型 API”，而在于把代码修复拆成了可建模、可搜索、可验证、可评估的算法问题。面试时你要主动把主线讲出来。"
    )
    add_number(doc, "代码理解：用 AST、CFG、Call Graph、Data-flow 把仓库转成结构化表示，减少纯文本 prompt 的不确定性。")
    add_number(doc, "缺陷定位：把失败测试、静态规则、图距离、数据流和语义信息融合成 FinalScore，输出 Top-k 可疑函数。")
    add_number(doc, "补丁搜索：不是只生成一次，而是产生多个候选，通过 Beam Search、去重、多样性重排和风险评分控制搜索。")
    add_number(doc, "执行验证：所有候选补丁进入 sandbox 跑 pytest，把“看起来对”变成“执行后通过”。")
    add_number(doc, "实验评估：用 benchmark、quality gate 和 ablation study 证明每个模块的贡献。")

    add_box(
        doc,
        [
            "一句话总结：这个项目的核心是“程序分析约束 LLM，搜索算法探索补丁空间，执行反馈关闭修复闭环，实验评估证明模块贡献”。",
        ],
        title="面试时可以背的核心句",
        fill=SOFT_GREEN,
    )


def add_boundaries(doc: Document) -> None:
    doc.add_heading("5. 简历中不要夸大的地方", level=1)
    add_table(
        doc,
        ["不要这样写", "原因", "更准确的写法"],
        [
            [
                "支持任意 GitHub 仓库一键自动修复并提交 PR",
                "当前项目完成的是受控 benchmark/template 驱动的闭环，不是任意真实 repo 的端到端 PR 自动化。",
                "支持 GitHub raw-source mutation benchmark 的可复现评估闭环。",
            ],
            [
                "大模型可以自动发现并修复所有 bug",
                "真实 bug 类型、测试覆盖和 ground truth 都不稳定，不能泛化夸大。",
                "在受控 benchmark 上验证多类缺陷定位与补丁修复能力。",
            ],
            [
                "Top-1=1.0000 代表真实项目上也能达到相同效果",
                "该指标来自受控 mutation benchmark，规则覆盖较强。",
                "在 62-case 受控 benchmark 上达到 Top-1 Localization 1.0000，并计划用真实 bug-fix commit benchmark 扩展泛化验证。",
            ],
        ],
        [2500, 3300, 3560],
    )


def add_interview_scripts(doc: Document) -> None:
    doc.add_heading("6. 面试讲解稿", level=1)
    add_box(
        doc,
        [
            "我做的是一个面向 Python 仓库的代码智能体，不是简单调用 LLM 修代码。系统先用 AST、CFG、Call Graph 和 Data-flow 构建程序图，再用 FinalScore 融合 SBFL、静态规则、图信号和语义信号做函数级缺陷定位。定位后只在 Top-k 可疑函数里生成补丁候选，并通过 Beam Search、pytest sandbox、execution feedback 和 reflection loop 做自动修复。最后我用 62-case GitHub raw-source mutation benchmark、quality gate 和 ablation study 验证定位、修复和各模块贡献。",
        ],
        title="30 秒版本",
        fill=SOFT_BLUE,
    )
    add_box(
        doc,
        [
            "这个项目分四层。第一层是代码理解层，负责解析 Python 仓库，抽取函数、类、调用、import、控制流和数据流，构建 Program Graph。第二层是缺陷定位层，核心是 FinalScore，它把失败测试覆盖、静态规则命中、调用图距离、数据流证据、语义相似度和风险信号融合起来，输出 Top-k suspicious functions。第三层是自动修复层，补丁生成器只在 Top-k 函数里生成候选 patch，并用 AST/scope validation 控制修改范围。第四层是搜索与评估层，用 Beam Search 在有限 sandbox budget 下选择候选，用 pytest 执行结果作为反馈，如果失败就进入 reflection loop 继续改。最终通过 benchmark、quality gate 和 ablation study 输出可复现报告。",
        ],
        title="2 分钟版本",
        fill=SOFT_BLUE,
    )


def add_qa(doc: Document) -> None:
    doc.add_heading("7. 高频面试问答", level=1)
    qa_rows = [
        [
            "Q1：它和直接调用 GPT 修代码有什么区别？",
            "直接调用 GPT 是把上下文丢给模型自由生成；这个项目先做程序分析和缺陷定位，只允许在 Top-k 可疑函数中产生小补丁，并用 sandbox 执行验证。LLM 是候选生成/语义判断的一部分，不是唯一决策者。",
        ],
        [
            "Q2：FinalScore 怎么设计？",
            "FinalScore 是多信号融合排序：SBFL 利用失败/通过测试覆盖，StaticRuleScore 捕捉明确 bug pattern，GraphScore 利用调用图和数据流接近度，SemanticScore/LLMScore 提供语义相关性，RiskScore 抑制高风险大范围修改。",
        ],
        [
            "Q3：为什么需要 Program Graph？",
            "单看源码文本很难知道失败如何传播。Program Graph 把函数、调用、控制流、数据流、import 和测试关系统一成图，方便做图距离、依赖传播、切片解释和跨函数定位。",
        ],
        [
            "Q4：为什么需要 Beam Search？",
            "补丁候选空间很大，sandbox 运行成本也高。Beam Search 保留多个高价值候选，并结合执行反馈、风险评分、多样性重排继续搜索，避免一次生成失败就结束。",
        ],
        [
            "Q5：reflection loop 做了什么？",
            "当补丁失败时，系统收集 traceback、失败断言、失败类型、已修改源码和图上下文，生成 refined patch。它把测试反馈转化成下一轮修复约束。",
        ],
        [
            "Q6：candidate deduplication 为什么重要？",
            "不同生成器可能产生等价补丁。如果不去重，sandbox budget 会被重复候选浪费。去重后可以把预算留给真正不同的修复假设。",
        ],
        [
            "Q7：diversity reranking 为什么有意义？",
            "如果前几个候选都来自同一个规则或同一种修复假设，一旦方向错了会连续失败。多样性重排让候选覆盖不同规则、目标函数、风险 bucket 和修复策略。",
        ],
        [
            "Q8：你如何证明模块有贡献？",
            "用 ablation study。完整系统跑一次，再关闭某个模块，比如静态规则、数据流、Beam Search、diversity reranking，对比 Top-1、MAP、Patch Success、Beam Success 等指标变化。",
        ],
        [
            "Q9：Top-1 指标很高，怎么解释？",
            "必须说明它来自受控 mutation benchmark，ground truth 明确且规则覆盖较强。它证明当前闭环有效，但不能夸大为任意真实世界 bug 都能达到同样效果。",
        ],
        [
            "Q10：后续怎么深化？",
            "优先做三件事：任意 GitHub repo 自动 benchmark onboarding、真实 bug-fix commit benchmark、Learning-to-Rank 版本的 FinalScore。",
        ],
    ]
    add_table(doc, ["问题", "回答口径"], qa_rows, [3000, 6360])


def add_deepening(doc: Document) -> None:
    doc.add_heading("8. 后续深化方向", level=1)
    add_table(
        doc,
        ["方向", "做什么", "为什么有价值"],
        [
            [
                "任意 GitHub repo benchmark onboarding",
                "输入 owner/repo/ref 后自动选择 Python 文件、生成 mutation template、补齐测试和 ground truth，并输出 benchmark report。",
                "把当前受控 benchmark 能力向更真实的仓库入口推进，提升项目完整度。",
            ],
            [
                "真实 bug-fix commit benchmark",
                "从真实修复提交中抽取 buggy/fixed 版本、patch、失败测试和 ground truth。",
                "验证系统在真实缺陷上的泛化能力，减少 mutation benchmark 偏差。",
            ],
            [
                "Learning-to-Rank FinalScore",
                "把 SBFL、GraphScore、StaticRuleScore、SemanticScore、RiskScore 等导出为特征，用 LightGBM Ranker/LambdaMART 学习排序。",
                "这是最能继续体现算法深度的方向，可和手工权重 FinalScore 做系统对比。",
            ],
            [
                "更强跨文件修复",
                "增强跨模块 call graph、data-flow、import alias、multi-patch repair 和 bundle search。",
                "让系统从单函数/小范围修复向真实仓库复杂修复靠近。",
            ],
        ],
        [2100, 3900, 3360],
    )


def add_learning_checklist(doc: Document) -> None:
    doc.add_heading("9. 准备面试时的学习顺序", level=1)
    add_number(doc, "先能解释 AST、CFG、Call Graph、Data-flow 分别解决什么问题。")
    add_number(doc, "再能画出整体流程：Repo Parser -> Program Graph -> FinalScore -> Patch Search -> Sandbox -> Reflection -> Report。")
    add_number(doc, "重点讲清 FinalScore：有哪些信号、为什么要融合、如何解释排序结果。")
    add_number(doc, "重点讲清 Beam Search：为什么需要搜索、如何用执行反馈调整候选。")
    add_number(doc, "最后准备边界说明：当前是受控 benchmark，不是任意 repo 一键真实 PR 修复。")

    add_box(
        doc,
        [
            "你面试时最应该主动说的一句话：",
            "“我把代码修复拆成了缺陷定位排序问题、补丁候选搜索问题和执行验证反馈问题，而不是把整个仓库直接交给 LLM。”",
        ],
        fill=SOFT_GREEN,
    )


def build_doc() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_final_resume_version(doc)
    add_versions(doc)
    add_role_mapping(doc)
    add_explanation(doc)
    add_boundaries(doc)
    add_interview_scripts(doc)
    add_qa(doc)
    add_deepening(doc)
    add_learning_checklist(doc)
    doc.core_properties.title = "Code Intelligence Agent System 简历写法与面试表达"
    doc.core_properties.subject = "算法向 LLM Agent 项目简历 bullet、技术栈、面试问答与后续深化"
    doc.core_properties.author = "Code Intelligence Agent"
    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
