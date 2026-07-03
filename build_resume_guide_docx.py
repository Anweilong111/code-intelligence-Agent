# -*- coding: utf-8 -*-
"""Build a resume-focused DOCX for the Code Intelligence Agent project."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

from build_project_report_docx import (
    APPLY_TABLE_GEOMETRY,
    CELL_MARGINS_DXA,
    CONTENT_WIDTH_DXA,
    LIGHT_BLUE_GRAY,
    NAVY,
    TABLE_INDENT_DXA,
    BLUE,
    DARK_BLUE,
    GRAY,
    configure_document,
    set_cell_shading,
    set_font_name,
)


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Code_Intelligence_Agent_Resume_Guide.docx"


def add_runs(paragraph, text: str, *, size: float = 10.8, bold: bool = False, color=None, font="Calibri") -> None:
    run = paragraph.add_run(text)
    set_font_name(run, font, "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str, *, size: float = 10.8, bold: bool = False, color=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    add_runs(p, text, size=size, bold=bold, color=color)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    add_runs(p, text, size=10.6)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    add_runs(p, text, size=10.6)
    return p


def add_code_box(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, "F4F6F9")
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, line, size=9.6, font="Consolas", color=RGBColor(0x22, 0x22, 0x22))
    APPLY_TABLE_GEOMETRY(
        table,
        [CONTENT_WIDTH_DXA],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120},
    )
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_font_name(run, "Calibri", "Microsoft YaHei")
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = NAVY

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font_name(run, "Calibri", "Microsoft YaHei")
                    run.font.size = Pt(9.4)
    APPLY_TABLE_GEOMETRY(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS_DXA,
    )
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, "简历包装与面试表达指南", size=12, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(title, "Code Intelligence Agent System", size=26, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(sub, "基于程序分析与 LLM 的代码缺陷定位与自动修复 Agent", size=14, color=DARK_BLUE)

    rows = [
        ["适合岗位", "算法工程师、LLM Agent 工程师、代码智能/软件工程智能方向、AI Infra/工具链方向"],
        ["简历定位", "程序分析 + 缺陷定位排序 + 补丁搜索 + sandbox 执行验证 + benchmark 评估"],
        ["核心指标", "62-case benchmark；Top-1 Localization 1.0000；MAP 1.0000；Patch Success Rate 1.0000；Beam Success Rate 0.9516"],
        ["生成日期", date.today().isoformat()],
    ]
    add_table(doc, ["项目", "内容"], rows, [2100, 7260])

    add_code_box(
        doc,
        [
            "简历表达原则：",
            "不要写成“调用大模型修代码”。",
            "要写成“程序图建模 + 多信号缺陷定位 + 多候选补丁搜索 + 执行验证闭环”。",
        ],
    )
    doc.add_page_break()


def build_doc() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("1. 简历上怎么定位这个项目", level=1)
    add_para(
        doc,
        "这个项目应该被包装成算法向 Agent 项目，而不是普通应用项目。简历重点不是“用了大模型”，而是你设计了一套可解释、可验证、可评估的代码智能体闭环。",
    )
    add_bullet(doc, "算法主线：AST / CFG / Call Graph / Data-flow / Program Graph 建模。")
    add_bullet(doc, "定位主线：静态规则、图传播、SBFL、LLM 语义评分融合成 FinalScore。")
    add_bullet(doc, "搜索主线：patch candidate generation、Beam Search、deduplication、diversity reranking、risk scoring。")
    add_bullet(doc, "验证主线：sandbox 中运行 pytest，用 execution feedback 和 reflection loop 迭代修复。")
    add_bullet(doc, "评估主线：benchmark、quality gate、ablation study、showcase report。")

    doc.add_heading("2. 项目名称推荐", level=1)
    add_code_box(
        doc,
        [
            "Code Intelligence Agent System：基于程序分析与 LLM 的代码缺陷定位与自动修复 Agent",
            "或：Program-Analysis-Guided LLM Agent for Bug Localization and Automated Repair",
        ],
    )

    doc.add_heading("3. 一句话项目描述", level=1)
    add_code_box(
        doc,
        [
            "设计并实现一个面向 Python 仓库的代码智能体系统，融合 AST/CFG/Call Graph/Data-flow 程序分析、",
            "多信号缺陷定位、补丁搜索与 sandbox 执行验证，实现从代码理解、Bug 定位到自动修复的闭环。",
        ],
    )

    doc.add_heading("4. 简历 Bullet：完整算法版", level=1)
    full_bullets = [
        "构建基于大语言模型与程序分析的代码智能体系统，支持 Repo Understanding、AST/CFG/Call Graph/Data-flow 建模、缺陷检测、函数级 Bug 定位、自动补丁生成与 sandbox 验证闭环。",
        "设计多信号缺陷定位算法 FinalScore，融合静态规则、调用图传播、数据流依赖、SBFL 可疑度与 LLM 语义评分，对可疑函数进行 Top-k 排序，并输出 attribution 解释定位原因。",
        "实现自动修复 pipeline：基于规则模板与 LLM 生成多候选 patch，通过 pytest sandbox 执行验证，并结合 execution feedback 与 reflection loop 迭代修复失败补丁。",
        "引入 Beam Search、candidate deduplication、diversity reranking 与 patch risk scoring，提升补丁搜索覆盖率，减少重复候选和高风险修改。",
        "构建 62 个 benchmark cases 的评估体系，覆盖跨函数数据流、程序切片、hard-case generation、quality gate 与 ablation study；在受控 benchmark 上实现 Top-1 Localization 1.0000、MAP 1.0000、Patch Success Rate 1.0000、Beam Success Rate 0.9516。",
    ]
    for item in full_bullets:
        add_bullet(doc, item)

    doc.add_heading("5. 简历 Bullet：精简版", level=1)
    add_code_box(
        doc,
        [
            "- 实现基于程序图与 LLM 的代码自动修复 Agent，融合 AST/CFG/Call Graph/Data-flow、SBFL 与语义评分完成函数级缺陷定位，并通过 Beam Search + sandbox pytest 验证实现多候选补丁搜索与自我修复闭环。",
            "- 构建 62-case benchmark 与 ablation 评估体系，在受控测试集上达到 Top-1 Localization 1.0000、MAP 1.0000、Patch Success Rate 1.0000。",
        ],
    )

    doc.add_heading("6. 不同岗位的写法", level=1)
    add_table(
        doc,
        ["投递方向", "简历强调重点", "推荐表达"],
        [
            ["算法工程师", "FinalScore、图建模、排序、消融实验", "突出多信号融合排序、Program Graph、SBFL、ablation study。"],
            ["LLM Agent 工程师", "Agent loop、工具调用、反馈闭环", "突出 patch generation、sandbox validation、reflection loop、LLM judge。"],
            ["软件工程/平台", "工程闭环、测试验证、报告体系", "突出 CLI、benchmark suite、pytest sandbox、质量门禁和可复现报告。"],
            ["AI Infra/工具链", "自动化、可观测、可扩展", "突出候选搜索、执行反馈、评估报告、失败分类和风险评分。"],
        ],
        [1700, 3100, 4560],
    )

    doc.add_heading("7. 技术栈写法", level=1)
    add_code_box(
        doc,
        [
            "Python, AST, CFG, Call Graph, Data-flow Analysis, Program Slicing, SBFL,",
            "LLM Agent, Patch Generation, Beam Search, pytest Sandbox, Benchmark, Ablation Study",
        ],
    )

    doc.add_heading("8. 简历中不要这样写", level=1)
    add_para(doc, "不要写得过浅：", bold=True, color=NAVY)
    add_code_box(doc, ["调用大模型自动修复代码"])
    add_para(doc, "也不要夸大边界：", bold=True, color=NAVY)
    add_code_box(doc, ["任意 GitHub 仓库一键自动修复，真实项目可直接上线"])
    add_para(doc, "更稳妥的表达是：", bold=True, color=NAVY)
    add_code_box(doc, ["在受控 benchmark/template 上完成从代码理解、缺陷定位、补丁生成到执行验证的自动修复闭环。"])

    doc.add_heading("9. 面试 30 秒讲解稿", level=1)
    add_code_box(
        doc,
        [
            "我做的是一个代码智能体系统，目标不是简单调用 LLM 修代码，而是把程序分析、缺陷定位、补丁搜索和执行验证串成闭环。",
            "系统会先解析 Python 仓库，构建 AST、CFG、Call Graph 和 Data-flow，再用 FinalScore 融合静态规则、图信号、SBFL 和语义评分定位可疑函数。",
            "定位后生成多个 patch candidate，通过 sandbox 跑 pytest 验证，并用 Beam Search 和 reflection loop 继续迭代。",
            "我还做了 benchmark、quality gate 和 ablation study，用指标证明每个模块的贡献。",
        ],
    )

    doc.add_heading("10. 面试 2 分钟讲解稿", level=1)
    add_para(
        doc,
        "这个项目分成四层：第一层是代码理解，第二层是缺陷定位，第三层是自动修复，第四层是实验评估。代码理解层负责解析仓库并构建 AST、CFG、Call Graph、Data-flow 和 Program Graph；缺陷定位层用 FinalScore 把静态规则、图传播、SBFL 和 LLM 语义评分融合起来，输出 Top-k suspicious functions 和 attribution；自动修复层根据定位结果生成多个补丁候选，并在 sandbox 中运行 pytest；如果失败，系统会读取执行反馈，通过 reflection loop 和 Beam Search 继续探索新的候选。最后，评估层通过 62-case benchmark、hard-case generation、quality gate 和 ablation study 验证系统效果。",
    )

    doc.add_heading("11. 高频面试问答", level=1)
    qa_rows = [
        ["Q1：和直接调用 GPT 修代码有什么区别？", "我的系统不是 LLM wrapper。LLM 只参与语义判断和补丁生成，前面有程序图建模和缺陷定位，后面有 sandbox 验证、搜索和评估。"],
        ["Q2：FinalScore 是什么？", "FinalScore 是函数级缺陷定位排序分数，融合静态规则、调用图/数据流图、SBFL 和 LLM 语义评分，用于决定优先修哪个函数。"],
        ["Q3：为什么要做 Program Graph？", "单文件 AST 只能看到局部结构，很多 bug 和跨函数调用、数据传递有关。Program Graph 能把函数、调用边、数据依赖和可疑信号统一起来。"],
        ["Q4：为什么需要 Beam Search？", "补丁空间不是单一路径，第一次生成的 patch 可能失败。Beam Search 保留多个高分候选，结合执行反馈继续探索，提高修复成功率。"],
        ["Q5：为什么需要 sandbox？", "自动修复不能只看 patch 文本是否合理，必须运行测试验证。sandbox 用隔离环境跑 pytest，避免污染原仓库，并把失败信息反馈给下一轮修复。"],
        ["Q6：Top-1 指标都是 1.0000 会不会太高？", "需要说明这是受控 benchmark/template 的结果，不代表任意真实仓库。项目当前重点是算法闭环和评估体系，不夸大泛化边界。"],
        ["Q7：项目最有算法含量的是哪里？", "多信号缺陷定位排序、程序图/切片建模、候选补丁搜索、diversity reranking、ablation-driven evaluation。"],
        ["Q8：后续怎么深化？", "一是任意 GitHub repo 自动 benchmark 化，二是真实 bug-fix commit benchmark，三是把 FinalScore 升级成 Learning-to-Rank。"],
    ]
    add_table(doc, ["问题", "回答口径"], qa_rows, [3100, 6260])

    doc.add_heading("12. 最终推荐放进简历的版本", level=1)
    add_code_box(
        doc,
        [
            "Code Intelligence Agent System：基于程序分析与 LLM 的代码缺陷定位与自动修复 Agent",
            "- 实现基于程序图与 LLM 的代码自动修复 Agent，融合 AST/CFG/Call Graph/Data-flow、SBFL 与语义评分完成函数级缺陷定位，并通过 Beam Search + sandbox pytest 验证实现多候选补丁搜索与自我修复闭环。",
            "- 构建 62-case benchmark 与 ablation 评估体系，在受控测试集上达到 Top-1 Localization 1.0000、MAP 1.0000、Patch Success Rate 1.0000。",
        ],
    )

    doc.core_properties.title = "Code Intelligence Agent System 简历包装与面试表达指南"
    doc.core_properties.subject = "简历 Bullet、技术栈、面试问答"
    doc.core_properties.author = "Code Intelligence Agent"
    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
