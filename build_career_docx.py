from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "docs" / "career"
ASSETS_DIR = OUTPUT_DIR / "assets"


DOCS = [
    (
        OUTPUT_DIR / "resume_agent_project_detailed.md",
        OUTPUT_DIR / "Code_Intelligence_Agent_Resume_Detailed_Guide.docx",
        "Code Intelligence Agent 简历写法详细指南",
        "用于简历项目经历、岗位适配、STAR 讲解和面试表述边界",
    ),
    (
        OUTPUT_DIR / "agent_project_study_interview_guide.md",
        OUTPUT_DIR / "Code_Intelligence_Agent_Study_Interview_Guide.docx",
        "Code Intelligence Agent 系统学习与面试准备指南",
        "从 0 理解 Agent 架构、算法模块、LLM 规划、安全门控、记忆系统和面试问答",
    ),
]


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ensure_diagram_assets()
    for markdown_path, docx_path, title, subtitle in DOCS:
        markdown = markdown_path.read_text(encoding="utf-8")
        document = build_document(markdown, title=title, subtitle=subtitle)
        document.save(docx_path)
        print(docx_path)


def build_document(markdown: str, *, title: str, subtitle: str) -> Document:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, title=title, subtitle=subtitle)
    add_document_note(
        doc,
        "阅读方式",
        "本文档按“是什么、为什么、怎么做、输出什么、面试怎么讲”的逻辑组织。建议先浏览目录式标题，再结合项目代码和 demo 报告逐章学习。",
    )
    add_markdown(doc, markdown)
    add_footer(doc, "Code Intelligence Agent")
    return doc


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_font(normal, "Calibri", 11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        set_font(style, "Calibri", size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    set_font(styles["List Bullet"], "Calibri", 11)
    styles["List Bullet"].paragraph_format.space_after = Pt(4)
    styles["List Bullet"].paragraph_format.line_spacing = 1.25
    set_font(styles["List Number"], "Calibri", 11)
    styles["List Number"].paragraph_format.space_after = Pt(4)
    styles["List Number"].paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    set_font(code_style, "Consolas", 9.5, color="1F2937")
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing = 1.05

    if "Callout" not in styles:
        callout = styles.add_style("Callout", 1)
    else:
        callout = styles["Callout"]
    set_font(callout, "Calibri", 10.5, color="1F3A5F")
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(6)
    callout.paragraph_format.line_spacing = 1.2


def set_font(style, font_name: str, size: float, *, color: str | None = None, bold: bool = False) -> None:
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title_block(doc: Document, *, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")

    add_horizontal_rule(doc)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_document_note(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.style = "Callout"
    add_formatted_runs(p, f"{label}：{body}")
    doc.add_paragraph()


def add_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        diagram = parse_diagram(line)
        if diagram:
            diagram_id, caption = diagram
            add_diagram(doc, diagram_id, caption)
            i += 1
            continue

        if is_table_start(lines, i):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        heading = parse_heading(line)
        if heading:
            level, text = heading
            if level == 1:
                if text.strip() == doc.paragraphs[0].text.strip():
                    i += 1
                    continue
                paragraph = doc.add_heading(text, level=1)
            elif level == 2:
                paragraph = doc.add_heading(text, level=2)
            else:
                paragraph = doc.add_heading(text, level=3)
            paragraph.paragraph_format.keep_with_next = True
            i += 1
            continue

        if line.startswith(">"):
            add_quote(doc, line.lstrip("> ").strip())
            i += 1
            continue

        bullet = parse_bullet(line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, bullet)
            i += 1
            continue

        numbered = parse_numbered(line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, numbered)
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            candidate = lines[i].rstrip()
            if not candidate.strip():
                break
            if (
                candidate.startswith("#")
                or candidate.startswith("```")
                or candidate.startswith("|")
                or parse_bullet(candidate)
                or parse_numbered(candidate)
                or candidate.startswith(">")
            ):
                break
            paragraph_lines.append(candidate)
            i += 1
        p = doc.add_paragraph()
        add_formatted_runs(p, " ".join(item.strip() for item in paragraph_lines))

    if in_code and code_lines:
        add_code_block(doc, "\n".join(code_lines))


def parse_heading(line: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.*)$", line)
    if not match:
        return None
    level = min(3, len(match.group(1)))
    return level, clean_inline(match.group(2).strip())


def parse_bullet(line: str) -> str | None:
    match = re.match(r"^\s*[-*]\s+(.*)$", line)
    return clean_inline(match.group(1).strip()) if match else None


def parse_numbered(line: str) -> str | None:
    match = re.match(r"^\s*\d+\.\s+(.*)$", line)
    return clean_inline(match.group(1).strip()) if match else None


def parse_diagram(line: str) -> tuple[str, str] | None:
    match = re.match(r"^\{\{DIAGRAM:([a-z0-9_-]+)(?:\|(.+))?\}\}$", line.strip())
    if not match:
        return None
    diagram_id = match.group(1)
    caption = match.group(2) or diagram_id.replace("_", " ")
    return diagram_id, clean_inline(caption)


def clean_inline(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text).strip()


def add_formatted_runs(paragraph, text: str) -> None:
    text = clean_inline(text)
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            add_run(paragraph, text[pos : match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = add_run(paragraph, token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string("1F2937")
        else:
            run = add_run(paragraph, token[2:-2])
            run.bold = True
        pos = match.end()
    if pos < len(text):
        add_run(paragraph, text[pos:])


def add_run(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return run


def add_code_block(doc: Document, text: str) -> None:
    for code_line in text.splitlines() or [""]:
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.left_indent = Inches(0.15)
        set_paragraph_shading(p, "F7F7F7")
        run = p.add_run(code_line if code_line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string("1F2937")


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Callout")
    p.paragraph_format.left_indent = Inches(0.18)
    set_paragraph_shading(p, "F4F6F9")
    add_formatted_runs(p, text)


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return lines[index].strip().startswith("|") and re.match(
        r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$",
        lines[index + 1],
    )


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = [split_table_row(line) for line in table_lines]
    if len(rows) < 2:
        return
    headers = rows[0]
    body_rows = rows[2:]
    col_count = len(headers)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.autofit = False

    widths = compute_table_widths(headers, body_rows)
    set_table_widths(table, widths)
    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = header_cells[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(clean_inline(text))
        run.bold = True
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)

    for row_values in body_rows:
        row_cells = table.add_row().cells
        for idx in range(col_count):
            text = row_values[idx] if idx < len(row_values) else ""
            cell = row_cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_formatted_runs(p, text)
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def compute_table_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    col_count = len(headers)
    if col_count == 1:
        return [9360]
    weights = []
    for idx in range(col_count):
        values = [headers[idx]] + [row[idx] for row in rows if idx < len(row)]
        max_len = max((len(value) for value in values), default=8)
        weights.append(max(8, min(max_len, 48)))
    total_weight = sum(weights)
    widths = [max(900, int(9360 * weight / total_weight)) for weight in weights]
    diff = 9360 - sum(widths)
    widths[-1] += diff
    return widths


def set_table_widths(table, widths: list[int]) -> None:
    table_width = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(table_width))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def ensure_diagram_assets() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    save_pipeline_diagram(ASSETS_DIR / "agent_pipeline.png")
    save_loop_diagram(ASSETS_DIR / "agent_loop.png")
    save_safety_diagram(ASSETS_DIR / "safety_gate.png")
    save_memory_diagram(ASSETS_DIR / "memory_chat.png")
    save_score_fusion_diagram(ASSETS_DIR / "score_fusion.png")
    save_program_graph_deep_diagram(ASSETS_DIR / "program_graph_deep.png")
    save_patch_repair_deep_diagram(ASSETS_DIR / "patch_repair_deep.png")
    save_controller_state_diagram(ASSETS_DIR / "controller_state.png")
    save_evaluation_metrics_diagram(ASSETS_DIR / "evaluation_metrics.png")


def add_diagram(doc: Document, diagram_id: str, caption: str) -> None:
    path = ASSETS_DIR / f"{diagram_id}.png"
    if not path.exists():
        add_quote(doc, f"图示占位：{caption}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.25))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    run = caption_paragraph.add_run(f"图示：{caption}")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")


def load_diagram_font(size: int, *, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str) -> None:
    x1, y1, x2, y2 = box
    max_chars = max(6, (x2 - x1) // 24)
    lines = wrap_diagram_text(text, max_chars)
    line_height = font.getbbox("测试Ag")[3] - font.getbbox("测试Ag")[1] + 8
    total_height = line_height * len(lines)
    y = y1 + ((y2 - y1) - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def wrap_diagram_text(text: str, max_chars: int) -> list[str]:
    lines: list[str] = []
    for part in text.splitlines() or [""]:
        if len(part) <= max_chars:
            lines.append(part)
            continue
        current = ""
        for char in part:
            current += char
            if len(current) >= max_chars:
                lines.append(current)
                current = ""
        if current:
            lines.append(current)
    return lines


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], *, fill: str = "#2E74B5") -> None:
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=5)
    if x2 >= x1:
        head = [(x2, y2), (x2 - 16, y2 - 9), (x2 - 16, y2 + 9)]
    else:
        head = [(x2, y2), (x2 + 16, y2 - 9), (x2 + 16, y2 + 9)]
    draw.polygon(head, fill=fill)


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font,
    *,
    fill: str = "#F4F6F9",
    outline: str = "#2E74B5",
    text_fill: str = "#0B2545",
) -> None:
    draw.rounded_rectangle(box, radius=24, fill=fill, outline=outline, width=4)
    draw_centered_text(draw, box, text, font, text_fill)


def save_pipeline_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 560), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(24, bold=True)
    small_font = load_diagram_font(20)
    draw.text((56, 36), "从 GitHub 仓库到智能分析报告的主链路", font=title_font, fill="#0B2545")
    nodes = [
        "输入仓库",
        "仓库发现",
        "AST/图建模",
        "缺陷信号",
        "测试诊断",
        "Agent 决策",
        "报告/记忆",
    ]
    x = 60
    y = 165
    width = 210
    gap = 40
    boxes = []
    for node in nodes:
        box = (x, y, x + width, y + 112)
        boxes.append(box)
        draw_box(draw, box, node, box_font)
        x += width + gap
    for left, right in zip(boxes, boxes[1:]):
        draw_arrow(draw, (left[2] + 8, (left[1] + left[3]) // 2), (right[0] - 8, (right[1] + right[3]) // 2))
    draw.rounded_rectangle((72, 360, 1728, 470), radius=18, fill="#E8EEF5", outline="#A9BBD0", width=2)
    draw.text(
        (100, 388),
        "关键点：不是把完整仓库一次性塞给 LLM，而是先结构化仓库、提取证据，再让 Controller/LLM 在安全边界内选择下一步。",
        font=small_font,
        fill="#1F3A5F",
    )
    img.save(path)


def save_loop_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(24, bold=True)
    small_font = load_diagram_font(20)
    draw.text((56, 36), "Agent Controller 闭环：Observe -> Plan -> Act -> Verify -> Reflect -> Replan", font=title_font, fill="#0B2545")
    boxes = {
        "Observe\n观察仓库状态": (90, 170, 360, 285),
        "Plan\n规划下一步": (510, 170, 780, 285),
        "Act\n执行动作": (930, 170, 1200, 285),
        "Verify\n验证结果": (1350, 170, 1620, 285),
        "Reflect\n失败反思": (930, 455, 1200, 570),
        "Replan\n重新规划": (510, 455, 780, 570),
    }
    for label, box in boxes.items():
        draw_box(draw, box, label, box_font)
    ordered = list(boxes.values())
    draw_arrow(draw, (360, 228), (510, 228))
    draw_arrow(draw, (780, 228), (930, 228))
    draw_arrow(draw, (1200, 228), (1350, 228))
    draw_arrow(draw, (1485, 285), (1065, 455))
    draw_arrow(draw, (930, 512), (780, 512))
    draw_arrow(draw, (510, 512), (225, 285))
    draw.rounded_rectangle((110, 645, 1690, 715), radius=18, fill="#F4F6F9", outline="#C8D4E3", width=2)
    draw.text((140, 665), "面试表达：Agent 的核心不是“用了大模型”，而是有状态、有动作、有验证、有失败后的重规划。", font=small_font, fill="#1F3A5F")
    img.save(path)


def save_safety_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 680), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(24, bold=True)
    small_font = load_diagram_font(20)
    draw.text((56, 36), "LLM 规划与安全门控：模型可以建议，但不能绕过执行边界", font=title_font, fill="#0B2545")
    top = [
        ((90, 170, 400, 300), "LLM Planner\n给出动作建议"),
        ((540, 170, 850, 300), "Action Registry\n动作白名单"),
        ((990, 170, 1300, 300), "Risk Policy\n风险约束"),
        ((1440, 170, 1710, 300), "Controller\n采纳或回退"),
    ]
    for box, text in top:
        draw_box(draw, box, text, box_font)
    for idx in range(len(top) - 1):
        draw_arrow(draw, (top[idx][0][2] + 8, 235), (top[idx + 1][0][0] - 8, 235))
    draw_box(draw, (280, 420, 760, 540), "通过：动作已注册、风险可控、预算允许", small_font, fill="#EEF7F0", outline="#2E7D32", text_fill="#1B5E20")
    draw_box(draw, (1030, 420, 1510, 540), "拒绝：动作未注册、越权、证据不足或预算超限", small_font, fill="#FFF4E5", outline="#B26A00", text_fill="#7A4D00")
    draw_arrow(draw, (855, 300), (520, 420), fill="#2E7D32")
    draw_arrow(draw, (1300, 300), (1270, 420), fill="#B26A00")
    draw.text((100, 600), "示例：LLM 推荐未注册动作时，安全门拒绝该建议，Controller 使用可审计的 fallback action。", font=small_font, fill="#1F3A5F")
    img.save(path)


def save_memory_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 660), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(24, bold=True)
    small_font = load_diagram_font(20)
    draw.text((56, 36), "多轮对话与记忆：围绕同一个 repo session 持续分析", font=title_font, fill="#0B2545")
    draw_box(draw, (90, 170, 400, 300), "用户消息\n继续修复/解释失败", box_font)
    draw_box(draw, (540, 170, 850, 300), "Session Memory\n历史状态", box_font)
    draw_box(draw, (990, 170, 1300, 300), "AgentController\n合并上下文", box_font)
    draw_box(draw, (1440, 170, 1710, 300), "下一步回复\n命令/报告/动作", box_font)
    draw_arrow(draw, (400, 235), (540, 235))
    draw_arrow(draw, (850, 235), (990, 235))
    draw_arrow(draw, (1300, 235), (1440, 235))
    draw_box(draw, (310, 420, 690, 535), "历史失败 patch\n用户约束\n测试摘要", small_font, fill="#E8EEF5")
    draw_box(draw, (1110, 420, 1490, 535), "避免重复失败\n保留边界\n可恢复会话", small_font, fill="#E8EEF5")
    draw_arrow(draw, (500, 420), (640, 300))
    draw_arrow(draw, (1300, 420), (1145, 300))
    draw.text((100, 590), "区别：这不是 ChatGPT 自由闲聊，而是围绕仓库分析 session 的连续任务记忆。", font=small_font, fill="#1F3A5F")
    img.save(path)


def save_score_fusion_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 800), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(23, bold=True)
    small_font = load_diagram_font(19)
    draw.text((56, 36), "FinalScore 多信号融合：从函数证据到 Top-k 排名", font=title_font, fill="#0B2545")
    nodes = [
        ((80, 155, 330, 265), "SBFL\n失败测试覆盖"),
        ((380, 155, 630, 265), "Graph\n图结构相关性"),
        ((680, 155, 930, 265), "Static\n静态规则置信度"),
        ((980, 155, 1230, 265), "Semantic\n语义相似度"),
        ((1280, 155, 1530, 265), "LLM\n模型评分"),
    ]
    for box, label in nodes:
        draw_box(draw, box, label, box_font)
        draw_arrow(draw, ((box[0] + box[2]) // 2, box[3] + 8), (900, 390))
    draw_box(draw, (680, 370, 1120, 500), "加权求和\n再减去 Risk 惩罚", box_font, fill="#E8EEF5")
    draw_box(draw, (1250, 370, 1580, 500), "Risk\n补丁风险惩罚", box_font, fill="#FFF4E5", outline="#B26A00", text_fill="#7A4D00")
    draw_arrow(draw, (1250, 435), (1128, 435), fill="#B26A00")
    draw_arrow(draw, (900, 508), (900, 600))
    draw_box(draw, (620, 600, 1180, 705), "FinalScore\n按分数降序得到 Top-k suspicious functions", box_font, fill="#EEF7F0", outline="#2E7D32", text_fill="#1B5E20")
    draw.text(
        (90, 740),
        "默认有覆盖证据：0.30*SBFL + 0.25*Graph + 0.15*Static + 0.10*Semantic + 0.15*LLM - 0.05*Risk。",
        font=small_font,
        fill="#1F3A5F",
    )
    img.save(path)


def save_program_graph_deep_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 820), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(22, bold=True)
    small_font = load_diagram_font(19)
    draw.text((56, 36), "Program Graph 深挖：节点、边与 GraphScore 证据来源", font=title_font, fill="#0B2545")
    left_nodes = [
        ((70, 150, 330, 245), "RepoParseResult\n文件/函数/类/import"),
        ((70, 300, 330, 395), "CallGraph\ncalls/awaits"),
        ((70, 450, 330, 545), "AST/CFG\n变量/控制流/基本块"),
    ]
    center_nodes = [
        ((520, 125, 795, 220), "file/class/function\ncontains"),
        ((520, 260, 795, 355), "calls/tested_by\nmodule_depends_on"),
        ((520, 395, 795, 490), "data_depends_on\narg/return flows"),
        ((520, 530, 795, 625), "controls/cfg_next\ncfg_branch/cfg_loop"),
    ]
    right_nodes = [
        ((1070, 180, 1350, 275), "proximity\nshortest path"),
        ((1070, 330, 1350, 425), "centrality\nPageRank"),
        ((1070, 480, 1350, 575), "caller impact\nmodule/async"),
    ]
    for box, text in left_nodes + center_nodes + right_nodes:
        draw_box(draw, box, text, box_font)
    draw_box(draw, (1450, 330, 1710, 450), "GraphScore\n多图信号融合", box_font, fill="#EEF7F0", outline="#2E7D32", text_fill="#1B5E20")
    for source_box, _ in left_nodes:
        draw_arrow(draw, (source_box[2] + 10, (source_box[1] + source_box[3]) // 2), (520, 370))
    for center_box, _ in center_nodes:
        draw_arrow(draw, (center_box[2] + 10, (center_box[1] + center_box[3]) // 2), (1070, 380))
    for right_box, _ in right_nodes:
        draw_arrow(draw, (right_box[2] + 10, (right_box[1] + right_box[3]) // 2), (1450, 390))
    draw.rounded_rectangle((90, 705, 1710, 775), radius=18, fill="#F4F6F9", outline="#C8D4E3", width=2)
    draw.text((120, 725), "面试重点：GraphScore 不是单一 centrality，而是 traceback、覆盖、最短路径、数据流、控制流、调用影响和风险的组合。", font=small_font, fill="#1F3A5F")
    img.save(path)


def save_patch_repair_deep_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 780), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(22, bold=True)
    small_font = load_diagram_font(19)
    draw.text((56, 36), "Patch 生成、Safety Gate、Sandbox 与 Reflection 闭环", font=title_font, fill="#0B2545")
    nodes = [
        ((70, 160, 340, 270), "Top-k 定位\n目标函数上下文"),
        ((460, 160, 730, 270), "Rule/LLM/Hybrid\n生成候选 patch"),
        ((850, 160, 1120, 270), "Safety Gate\nAST/Scope/签名/大小"),
        ((1240, 160, 1510, 270), "Sandbox\n复制仓库+pytest"),
    ]
    for box, text in nodes:
        draw_box(draw, box, text, box_font)
    for left, right in zip(nodes, nodes[1:]):
        draw_arrow(draw, (left[0][2] + 8, 215), (right[0][0] - 8, 215))
    draw_box(draw, (1240, 430, 1510, 540), "成功\n记录 verified repair", box_font, fill="#EEF7F0", outline="#2E7D32", text_fill="#1B5E20")
    draw_box(draw, (850, 430, 1120, 540), "失败反馈\nstdout/stderr/traceback", box_font, fill="#FFF4E5", outline="#B26A00", text_fill="#7A4D00")
    draw_box(draw, (460, 430, 730, 540), "Reflection\n分类失败+避免重复", box_font, fill="#F4F6F9")
    draw_arrow(draw, (1375, 270), (1375, 430), fill="#2E7D32")
    draw_arrow(draw, (1240, 485), (1120, 485), fill="#B26A00")
    draw_arrow(draw, (850, 485), (730, 485), fill="#B26A00")
    draw_arrow(draw, (595, 430), (595, 270), fill="#B26A00")
    draw.rounded_rectangle((90, 650, 1710, 720), radius=18, fill="#E8EEF5", outline="#A9BBD0", width=2)
    draw.text((120, 670), "关键边界：LLM 只能生成候选；Safety Gate 和 Sandbox 才决定候选是否可执行、是否可采纳。", font=small_font, fill="#1F3A5F")
    img.save(path)


def save_controller_state_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 800), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(21, bold=True)
    small_font = load_diagram_font(19)
    draw.text((56, 36), "AgentController 深挖：Stage/Blocker -> Action -> Verify -> Replan", font=title_font, fill="#0B2545")
    stages = [
        ((70, 155, 330, 255), "source_import_blocked\n调整源码/缓存/token"),
        ((70, 300, 330, 400), "phase1\n仓库理解/测试发现"),
        ((70, 445, 330, 545), "phase2\n静态/动态图定位"),
        ((70, 590, 330, 690), "phase3\npatch validation"),
    ]
    middle = [
        ((560, 190, 850, 290), "Action Registry\n动作白名单+别名"),
        ((560, 395, 850, 495), "LLM Planner\nadvisory only"),
        ((560, 600, 850, 700), "Memory Context\nsession/repo/repair"),
    ]
    for box, text in stages + middle:
        draw_box(draw, box, text, box_font)
    draw_box(draw, (1110, 315, 1410, 435), "Safety Gate\n未注册/越权拒绝", box_font, fill="#FFF4E5", outline="#B26A00", text_fill="#7A4D00")
    draw_box(draw, (1510, 315, 1725, 435), "Selected Action\n可执行或 blocker", box_font, fill="#EEF7F0", outline="#2E7D32", text_fill="#1B5E20")
    for stage_box, _ in stages:
        draw_arrow(draw, (stage_box[2] + 10, (stage_box[1] + stage_box[3]) // 2), (560, 395))
    for mid_box, _ in middle:
        draw_arrow(draw, (mid_box[2] + 10, (mid_box[1] + mid_box[3]) // 2), (1110, 375))
    draw_arrow(draw, (1410, 375), (1510, 375))
    draw.rounded_rectangle((420, 80, 1710, 130), radius=16, fill="#F4F6F9", outline="#C8D4E3", width=2)
    draw.text((445, 94), "规则 Controller 始终保留最终执行权；LLM 推荐只能匹配注册 action 或作为 advisory 记录。", font=small_font, fill="#1F3A5F")
    img.save(path)


def save_evaluation_metrics_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(36, bold=True)
    box_font = load_diagram_font(22, bold=True)
    small_font = load_diagram_font(19)
    draw.text((56, 36), "评估体系深挖：定位指标、权重搜索、Ablation 与置信度校准", font=title_font, fill="#0B2545")
    nodes = [
        ((80, 160, 360, 270), "Benchmark Cases\nranked + ground truth"),
        ((500, 160, 780, 270), "Localization Metrics\nTop-k/MRR/MAP/EXAM"),
        ((920, 160, 1200, 270), "Weight Search\nvalidation score"),
        ((1340, 160, 1620, 270), "Robustness\nholdout gap penalty"),
    ]
    for box, text in nodes:
        draw_box(draw, box, text, box_font)
    for left, right in zip(nodes, nodes[1:]):
        draw_arrow(draw, (left[0][2] + 8, 215), (right[0][0] - 8, 215))
    lower = [
        ((300, 450, 620, 560), "Ablation\nwithout_* variants"),
        ((740, 450, 1060, 560), "Calibration\nBrier/ECE"),
        ((1180, 450, 1500, 560), "面试结论\n不是拍脑袋权重"),
    ]
    for box, text in lower:
        draw_box(draw, box, text, box_font, fill="#E8EEF5")
    draw_arrow(draw, (640, 270), (460, 450))
    draw_arrow(draw, (1060, 270), (900, 450))
    draw_arrow(draw, (1200, 270), (1340, 450))
    draw.rounded_rectangle((90, 650, 1710, 715), radius=18, fill="#F4F6F9", outline="#C8D4E3", width=2)
    draw.text((120, 668), "关键表达：当前权重是可解释启发式，并通过 metrics、weight_search、ablation 和 calibration 做验证与改进。", font=small_font, fill="#1F3A5F")
    img.save(path)


def add_footer(doc: Document, label: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.text = label
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("666666")


if __name__ == "__main__":
    main()
